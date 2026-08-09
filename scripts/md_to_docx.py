#!/usr/bin/env python3
"""
Convert paper/methodology.md to a .docx file with:
  - Times New Roman, 12pt body text
  - Proper heading hierarchy
  - Readable math (converted from LaTeX to Unicode where possible)
  - Tables preserved
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)

# ── LaTeX → Unicode math conversion helpers ──
def convert_math(text):
    """Convert common LaTeX math to readable Unicode."""
    # Greek letters
    greek = {
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
        r'\theta': 'θ', r'\lambda': 'λ', r'\mu': 'μ', r'\sigma': 'σ',
        r'\tau': 'τ', r'\phi': 'φ', r'\omega': 'ω', r'\rho': 'ρ',
        r'\mathcal': '', r'\mathbf': '', r'\boldsymbol': '', r'\mathrm': '',
        r'\mathbb': '', r'\odot': '⊙', r'\cdot': '·', r'\times': '×',
        r'\in': '∈', r'\cup': '∪', r'\bigcup': '∪', r'\sum': 'Σ',
        r'\sqrt': '√', r'\to': '→', r'\mapsto': '→', r'\ge': '≥',
        r'\le': '≤', r'\neq': '≠', r'\approx': '≈', r'\infty': '∞',
        r'\odot': '⊙', r'\odot': '⊙',
    }
    for k, v in greek.items():
        text = text.replace(k, v)
    # Subscripts: _{...} and _x
    text = re.sub(r'_\{([^}]*)\}', r'[\1]', text)
    text = re.sub(r'_([a-zA-Z0-9])', r'[\1]', text)
    # Superscripts: ^{...} and ^x
    text = re.sub(r'\^\{([^}]*)\}', r'^(\1)', text)
    text = re.sub(r'\^([a-zA-Z0-9])', r'^(\1)', text)
    # Remove remaining braces
    text = text.replace('{', '').replace('}', '')
    # Fractions
    text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)
    # Operators
    text = text.replace(r'\operatorname{clamp}', 'clamp')
    text = text.replace(r'\operatorname{BN}', 'BN')
    text = text.replace(r'\operatorname{softmax}', 'softmax')
    text = text.replace(r'\exp', 'exp')
    text = text.replace(r'\log', 'log')
    text = text.replace(r'\min', 'min')
    text = text.replace(r'\max', 'max')
    text = text.replace(r'\big(', '(').replace(r'\big)', ')')
    text = text.replace(r'\,', ' ').replace(r'\;', ' ')
    text = text.replace(r'\ ', ' ')
    return text.strip()


def process_inline_math(text):
    """Convert $...$ inline math to Unicode."""
    def repl(m):
        return convert_math(m.group(1))
    return re.sub(r'\$([^$]+)\$', repl, text)


def process_display_math(text):
    """Convert $$...$$ display math to a readable Unicode block."""
    def repl(m):
        return convert_math(m.group(1))
    return re.sub(r'\$\$([^$]+)\$\$', repl, text)


def add_math_paragraph(doc, math_text, is_centered=True):
    """Add a display-math paragraph."""
    p = doc.add_paragraph()
    if is_centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(convert_math(math_text))
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE
    run.italic = True
    return p


def set_cell_text(cell, text, bold=False):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(process_inline_math(text))
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.bold = bold


def main():
    src = os.path.join(os.path.dirname(__file__), "..", "paper", "methodology.md")
    src = os.path.abspath(src)
    out = os.path.join(os.path.dirname(src), "methodology.docx")

    with open(src, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Display math block
        if stripped.startswith("$$"):
            math_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("$$"):
                math_lines.append(lines[i].strip())
                i += 1
            i += 1  # skip closing $$
            add_math_paragraph(doc, " ".join(math_lines))
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(process_inline_math(stripped[2:]), level=1)
            for run in h.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(16)
            i += 1
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(process_inline_math(stripped[3:]), level=2)
            for run in h.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            i += 1
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(process_inline_math(stripped[4:]), level=3)
            for run in h.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(13)
            i += 1
            continue

        # Table detection
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r'^\s*\|[\s:|-]+\|', lines[i+1]):
            # Parse table
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header and separator
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            # Create table
            ncols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, htext in enumerate(header):
                set_cell_text(table.rows[0].cells[j], htext, bold=True)
            for ri, row in enumerate(rows):
                for j in range(min(ncols, len(row))):
                    set_cell_text(table.rows[ri+1].cells[j], row[j])
            doc.add_paragraph()  # spacing after table
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(process_inline_math(stripped[2:]))
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(process_inline_math(re.sub(r'^\d+\.\s', '', stripped)))
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            i += 1
            continue

        # Bold paragraph (e.g., **G1 — ...**)
        if stripped.startswith("**") and "**" in stripped[2:]:
            p = doc.add_paragraph()
            # Split bold and normal parts
            parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(process_inline_math(part[2:-2]))
                    run.bold = True
                elif part:
                    run = p.add_run(process_inline_math(part))
                if run:
                    run.font.name = FONT_NAME
                    run.font.size = BODY_SIZE
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        text = process_inline_math(stripped)
        # Handle inline bold within normal text
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                run = p.add_run(part)
            if run:
                run.font.name = FONT_NAME
                run.font.size = BODY_SIZE
        i += 1

    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
